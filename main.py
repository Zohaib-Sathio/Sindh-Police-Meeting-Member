"""
Sindh Police AI Meeting Member - Main Application
Real-time AI meeting participation system for Sindh Police Department
"""

import os
import sys

# Set UTF-8 encoding for Windows console
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
import json
import base64
import asyncio
import websockets
import uuid
import time
import io
import hashlib
import re
from fastapi import FastAPI, WebSocket, Request, HTTPException, Body, UploadFile, File
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.websockets import WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from datetime import datetime as dt, timedelta, timezone
from typing import List
import jwt
from dotenv import load_dotenv
from pydub import AudioSegment
import audioop
from contextlib import suppress

from prompts import function_call_tools, build_system_message
from tools import (
    start_meeting_session,
    end_meeting_session,
    cast_vote,
    add_motion,
    get_meeting_status,
    get_vote_history,
    add_transcript_entry,
    get_transcript,
    request_regulatory_context,
    meeting_sessions,
    vote_history
)

load_dotenv(override=True)

# --- Configuration ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PORT = int(os.getenv("PORT", 8000))

VOICE = 'sage'

LOG_EVENT_TYPES = [
    'error', 'response.content.done', 'rate_limits.updated',
    'input_audio_buffer.committed', 'session.created'
]

SHOW_TIMING_MATH = False
call_recordings = {}

app = FastAPI(
    title="Sindh Police AI Meeting Member",
    description="AI-powered meeting participation system for Sindh Police Department",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000", "http://127.0.0.1:5173", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# JWT Configuration
JWT_SECRET_KEY = os.getenv("JWT_SECRET_KEY", "sindh-police-ai-meeting-member-secret-key-2024")
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 24

# User roles for Sindh Police
USERS_DB = {
    "secretary": {
        "username": "secretary",
        "password": "secretary123",
        "full_name": "Board Secretary",
        "role": "secretary"
    },
    "admin": {
        "username": "admin",
        "password": "admin123",
        "full_name": "System Administrator",
        "role": "admin"
    },
    "admin@sindhpolice.gov.pk": {
        "username": "admin@sindhpolice.gov.pk",
        "password": "sindhpolice@ai",
        "full_name": "System Administrator",
        "role": "admin"
    },
    "observer": {
        "username": "observer",
        "password": "observer123",
        "full_name": "Board Observer",
        "role": "observer"
    }
}

from fastapi.staticfiles import StaticFiles
app.mount("/client", StaticFiles(directory="static", html=True), name="client")

CHANNELS = 1
RATE = 8000

# Meeting metadata storage
meeting_metadata: dict[str, dict] = {}
active_meeting_id: str = None


@app.get("/", response_class=HTMLResponse)
async def index_page():
    """Serve the Live Meeting Dashboard (Home)"""
    with open("static/voice-client.html", "r", encoding="utf-8") as f:
        html_content = f.read()
    return html_content


# =============================================================================
# AUTHENTICATION ENDPOINTS
# =============================================================================

def create_jwt_token(username: str, full_name: str, role: str) -> str:
    """Create a JWT token for the user"""
    now = dt.now(timezone.utc)
    payload = {
        "username": username,
        "full_name": full_name,
        "role": role,
        "exp": now + timedelta(hours=JWT_EXPIRATION_HOURS),
        "iat": now
    }
    token = jwt.encode(payload, JWT_SECRET_KEY, algorithm=JWT_ALGORITHM)
    return token


def verify_jwt_token(token: str) -> dict:
    """Verify and decode JWT token"""
    try:
        payload = jwt.decode(token, JWT_SECRET_KEY, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token has expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


def get_token_from_request(request: Request) -> str:
    """Extract JWT token from Authorization header"""
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid authorization header")
    return auth_header.replace("Bearer ", "")


@app.post("/auth/login")
async def login(credentials: dict = Body(...)):
    """Authenticate user with username and password"""
    username = credentials.get("username", "").strip()
    password = credentials.get("password", "")
    
    if username in USERS_DB:
        user = USERS_DB[username]
        if user["password"] == password:
            token = create_jwt_token(username, user["full_name"], user["role"])
            
            return {
                "success": True,
                "message": "Login successful",
                "token": token,
                "user": {
                    "username": username,
                    "full_name": user["full_name"],
                    "role": user["role"]
                }
            }
    
    raise HTTPException(status_code=401, detail="Invalid username or password")


# =============================================================================
# MEETING MANAGEMENT ENDPOINTS
# =============================================================================

@app.post("/api/meeting/start")
async def api_start_meeting(request: Request, payload: dict = Body(...)):
    """Start a new board meeting session"""
    global active_meeting_id
    
    token = get_token_from_request(request)
    user_data = verify_jwt_token(token)
    
    # Only secretary and admin can start meetings
    if user_data.get("role") not in ["secretary", "admin"]:
        raise HTTPException(status_code=403, detail="Only Secretary or Admin can start meetings")
    
    meeting_id = payload.get("meeting_id", f"MEETING-{dt.now().strftime('%Y%m%d%H%M%S')}")
    agenda = payload.get("agenda", "")
    
    result = start_meeting_session(meeting_id, agenda)
    active_meeting_id = meeting_id
    
    meeting_metadata[meeting_id] = {
        "started_by": user_data["username"],
        "agenda": agenda
    }
    
    return result


@app.post("/api/meeting/end")
async def api_end_meeting(request: Request, payload: dict = Body(...)):
    """End the current board meeting session and automatically generate meeting notes"""
    global active_meeting_id
    
    token = get_token_from_request(request)
    user_data = verify_jwt_token(token)
    
    if user_data.get("role") not in ["secretary", "admin"]:
        raise HTTPException(status_code=403, detail="Only Secretary or Admin can end meetings")
    
    meeting_id = payload.get("meeting_id", active_meeting_id)
    
    if not meeting_id:
        raise HTTPException(status_code=400, detail="No active meeting to end")
    
    result = end_meeting_session(meeting_id)
    
    if meeting_id == active_meeting_id:
        active_meeting_id = None
    
    # Automatically generate meeting notes
    meeting_notes = None
    try:
        # Get transcript
        transcript = get_transcript(meeting_id)
        if transcript:
            # Get meeting details
            meeting_info = {}
            if meeting_id in meeting_sessions:
                meeting_info = meeting_sessions[meeting_id]
            
            # Format transcript for AI processing
            transcript_text = ""
            for entry in transcript:
                speaker = entry.get("speaker", "Unknown")
                text = entry.get("text", "")
                timestamp = entry.get("timestamp", "")
                transcript_text += f"[{timestamp}] {speaker}: {text}\n"
            
            # Get votes for context
            votes = get_vote_history(meeting_id)
            votes_text = ""
            if votes:
                votes_text = "\n\nVOTES CAST:\n"
                for vote in votes:
                    votes_text += f"- Motion: {vote.get('motion_description', 'N/A')}\n"
                    votes_text += f"  Vote: {vote.get('vote', 'N/A')}\n"
                    votes_text += f"  Reasoning: {vote.get('reasoning', 'N/A')}\n\n"
            
            # Generate meeting notes using OpenAI
            import openai
            client = openai.OpenAI(api_key=OPENAI_API_KEY)
            
            meeting_date = meeting_info.get("start_time", "")
            if meeting_date:
                try:
                    from zoneinfo import ZoneInfo
                    karachi_tz = ZoneInfo("Asia/Karachi")
                    dt_obj = dt.fromisoformat(meeting_date.replace('Z', '+00:00'))
                    dt_obj = dt_obj.astimezone(karachi_tz)
                    meeting_date = dt_obj.strftime("%B %d, %Y at %I:%M %p")
                except:
                    pass
            
            notes_prompt = f"""You are a professional meeting secretary for the Sindh Police Department Meeting.

Generate comprehensive, well-formatted meeting notes from the following transcript.

MEETING INFORMATION:
- Meeting ID: {meeting_id}
- Date: {meeting_date}
- Status: {meeting_info.get('status', 'Unknown')}

TRANSCRIPT:
{transcript_text}

{votes_text}

Please generate professional meeting notes in the following format:

# SINDH POLICE DEPARTMENT MEETING NOTES

**Meeting ID:** {meeting_id}
**Date:** {meeting_date}
**Status:** {meeting_info.get('status', 'Unknown')}

## Meeting Summary
[Provide a concise 2-3 sentence summary of the meeting]

## Discussion Points
[Organize the discussion into clear topics/agenda items]

## User Insights and Discussion Summary
[Provide a comprehensive summary of the board members' (users') insights, perspectives, concerns, and key points raised during the discussion. Focus on:
- Main concerns and viewpoints expressed by board members
- Different perspectives or opinions shared
- Important insights or observations made
- Questions raised and clarifications sought
- Areas of agreement or disagreement among members
- Any suggestions or recommendations put forward by board members]

## Key Decisions
[List any decisions made during the meeting]

## Votes Cast
[Include all votes with motion, vote (FOR/AGAINST/ABSTAIN), and reasoning]

## Action Items
[Extract any action items or follow-ups mentioned]

## Attendees
[Based on speakers in transcript, list attendees]

## Next Steps
[Any next steps or future agenda items mentioned]

---

*Generated automatically by Sindh Police AI Meeting Member System*

Make the notes professional, clear, and well-organized. Use proper formatting with markdown."""
            
            notes_response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a professional meeting secretary. Generate clear, well-formatted meeting notes from transcripts."},
                    {"role": "user", "content": notes_prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            
            meeting_notes = notes_response.choices[0].message.content.strip()
            print(f"✅ Meeting notes generated for {meeting_id}")
        else:
            print(f"⚠️ No transcript found for {meeting_id}, skipping notes generation")
    except Exception as e:
        print(f"⚠️ Error generating meeting notes: {e}")
        import traceback
        traceback.print_exc()
        # Don't fail the meeting end if notes generation fails
        meeting_notes = None
    
    # Add meeting notes to result
    if meeting_notes:
        result["meeting_notes"] = meeting_notes
        result["notes_generated"] = True
    else:
        result["notes_generated"] = False
    
    # Add full meeting minutes data for the frontend popup
    if meeting_id in meeting_sessions:
        session = meeting_sessions[meeting_id]
        result["minutes"] = {
            "meeting_id": meeting_id,
            "start_time": session.get("start_time", ""),
            "end_time": session.get("end_time", ""),
            "agenda": session.get("agenda", ""),
            "transcript": session.get("transcript", []),
            "votes": session.get("votes", []),
            "motions": session.get("motions", []),
        }
    
    return result


@app.get("/api/meeting/status")
async def api_meeting_status(request: Request):
    """Get current meeting status"""
    token = get_token_from_request(request)
    verify_jwt_token(token)
    
    if not active_meeting_id:
        return {"active": False, "message": "No active meeting"}
    
    return {
        "active": True,
        "meeting_id": active_meeting_id,
        **get_meeting_status(active_meeting_id)
    }


@app.post("/api/motion/add")
async def api_add_motion(request: Request, payload: dict = Body(...)):
    """Add a new motion for voting"""
    token = get_token_from_request(request)
    user_data = verify_jwt_token(token)
    
    if user_data.get("role") not in ["secretary", "admin"]:
        raise HTTPException(status_code=403, detail="Only Secretary or Admin can add motions")
    
    if not active_meeting_id:
        raise HTTPException(status_code=400, detail="No active meeting")
    
    motion_text = payload.get("motion_text", "")
    if not motion_text:
        raise HTTPException(status_code=400, detail="Motion text is required")
    
    result = add_motion(active_meeting_id, motion_text, user_data["username"])
    return result


@app.get("/api/votes/history")
async def api_vote_history(request: Request, meeting_id: str = None):
    """Get vote history"""
    token = get_token_from_request(request)
    verify_jwt_token(token)
    
    votes = get_vote_history(meeting_id)
    return {"votes": votes}


@app.get("/api/transcript/{meeting_id}")
async def api_get_transcript(meeting_id: str, request: Request):
    """Get meeting transcript"""
    token = get_token_from_request(request)
    verify_jwt_token(token)
    
    transcript = get_transcript(meeting_id)
    return {"transcript": transcript}


@app.post("/api/meeting/notes/generate")
async def generate_meeting_notes(request: Request, payload: dict = Body(...)):
    """Generate formatted meeting notes from transcript using AI"""
    token = get_token_from_request(request)
    user_data = verify_jwt_token(token)
    
    if user_data.get("role") not in ["secretary", "admin"]:
        raise HTTPException(status_code=403, detail="Admin or Secretary access required")
    
    meeting_id = payload.get("meeting_id")
    if not meeting_id:
        raise HTTPException(status_code=400, detail="meeting_id is required")
    
    try:
        import openai
        
        # Get transcript
        transcript = get_transcript(meeting_id)
        if not transcript:
            raise HTTPException(status_code=404, detail="No transcript found for this meeting")
        
        # Get meeting details
        meeting_info = {}
        if meeting_id in meeting_sessions:
            meeting_info = meeting_sessions[meeting_id]
        
        # Format transcript for AI processing
        transcript_text = ""
        for entry in transcript:
            speaker = entry.get("speaker", "Unknown")
            text = entry.get("text", "")
            timestamp = entry.get("timestamp", "")
            transcript_text += f"[{timestamp}] {speaker}: {text}\n"
        
        # Get votes for context
        votes = get_vote_history(meeting_id)
        votes_text = ""
        if votes:
            votes_text = "\n\nVOTES CAST:\n"
            for vote in votes:
                votes_text += f"- Motion: {vote.get('motion_description', 'N/A')}\n"
                votes_text += f"  Vote: {vote.get('vote', 'N/A')}\n"
                votes_text += f"  Reasoning: {vote.get('reasoning', 'N/A')}\n\n"
        
        # Generate meeting notes using OpenAI
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        
        meeting_date = meeting_info.get("start_time", "")
        if meeting_date:
            try:
                from zoneinfo import ZoneInfo
                karachi_tz = ZoneInfo("Asia/Karachi")
                dt_obj = dt.fromisoformat(meeting_date.replace('Z', '+00:00'))
                dt_obj = dt_obj.astimezone(karachi_tz)
                meeting_date = dt_obj.strftime("%B %d, %Y at %I:%M %p")
            except:
                pass
        
        notes_prompt = f"""You are a professional meeting secretary for the Sindh Police Department Meeting.

Generate comprehensive, well-formatted meeting notes from the following transcript.

MEETING INFORMATION:
- Meeting ID: {meeting_id}
- Date: {meeting_date}
- Status: {meeting_info.get('status', 'Unknown')}

TRANSCRIPT:
{transcript_text}

{votes_text}

Please generate professional meeting notes in the following format:

# SINDH POLICE DEPARTMENT MEETING NOTES

**Meeting ID:** {meeting_id}
**Date:** {meeting_date}
**Status:** {meeting_info.get('status', 'Unknown')}

## Meeting Summary
[Provide a concise 2-3 sentence summary of the meeting]

## Discussion Points
[Organize the discussion into clear topics/agenda items]

## User Insights and Discussion Summary
[Provide a comprehensive summary of the board members' (users') insights, perspectives, concerns, and key points raised during the discussion. Focus on:
- Main concerns and viewpoints expressed by board members
- Different perspectives or opinions shared
- Important insights or observations made
- Questions raised and clarifications sought
- Areas of agreement or disagreement among members
- Any suggestions or recommendations put forward by board members]

## Key Decisions
[List any decisions made during the meeting]

## Votes Cast
[Include all votes with motion, vote (FOR/AGAINST/ABSTAIN), and reasoning]

## Action Items
[Extract any action items or follow-ups mentioned]

## Attendees
[Based on speakers in transcript, list attendees]

## Next Steps
[Any next steps or future agenda items mentioned]

---

*Generated automatically by Sindh Police AI Meeting Member System*

Make the notes professional, clear, and well-organized. Use proper formatting with markdown."""
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a professional meeting secretary. Generate clear, well-formatted meeting notes from transcripts."},
                {"role": "user", "content": notes_prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        meeting_notes = response.choices[0].message.content.strip()
        
        return {
            "success": True,
            "meeting_id": meeting_id,
            "notes": meeting_notes,
            "generated_at": dt.now(timezone.utc).isoformat()
        }
        
    except Exception as e:
        print(f"❌ Meeting notes generation error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate meeting notes: {str(e)}")


@app.post("/api/meeting/notes/download-docx")
async def download_meeting_notes_docx(request: Request):
    """Convert meeting notes markdown to DOCX format and return as download"""
    try:
        # Parse body manually to get better error messages
        try:
            body = await request.json()
        except Exception as e:
            print(f"❌ JSON parsing error: {e}")
            raise HTTPException(status_code=400, detail=f"Invalid JSON in request body: {str(e)}")
        
        # Debug: Log raw request
        print(f"📄 DOCX Download Request received")
        print(f"   Request method: {request.method}")
        print(f"   Content-Type: {request.headers.get('content-type', 'N/A')}")
        
        token = get_token_from_request(request)
        user_data = verify_jwt_token(token)
        
        if user_data.get("role") not in ["secretary", "admin"]:
            raise HTTPException(status_code=403, detail="Admin or Secretary access required")
        
        # Debug: Log body
        print(f"   Body type: {type(body)}")
        print(f"   Body keys: {list(body.keys()) if isinstance(body, dict) else 'NOT A DICT'}")
        
        meeting_id = body.get("meeting_id") if isinstance(body, dict) else None
        notes = body.get("notes", "") if isinstance(body, dict) else ""
        
        # Debug logging
        print(f"   Meeting ID: {meeting_id}")
        print(f"   Notes length: {len(notes) if notes else 0}")
        print(f"   Notes preview: {notes[:100] if notes else 'EMPTY'}...")
        
        if not meeting_id:
            print("❌ Missing meeting_id")
            raise HTTPException(status_code=400, detail="meeting_id is required")
        
        if not notes or not notes.strip():
            print("❌ Missing or empty notes")
            raise HTTPException(status_code=400, detail="notes content is required and cannot be empty")
        
        # Log for debugging
        print(f"✅ Generating DOCX for meeting: {meeting_id}, notes length: {len(notes)}")
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error in DOCX endpoint validation: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"Invalid request: {str(e)}")
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        from io import BytesIO
        
        # Create a new Document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Parse markdown and convert to DOCX
        lines = notes.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Handle headers
            if line.startswith('# '):
                # H1
                p = doc.add_heading(line[2:].strip(), level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                # H2
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                # H3
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('**') and line.endswith('**'):
                # Bold text (metadata)
                text = line.replace('**', '')
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
            elif line.startswith('- '):
                # Bullet point
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif line.startswith('---'):
                # Horizontal rule
                doc.add_paragraph('_' * 50)
            elif line.strip():
                # Regular paragraph
                # Handle inline markdown (bold, italic)
                paragraph = doc.add_paragraph()
                add_formatted_text(paragraph, line)
            else:
                # Empty line
                doc.add_paragraph()
            
            i += 1
        
        # Save to BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        from fastapi.responses import Response
        
        return Response(
            content=docx_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="meeting-notes-{meeting_id}.docx"'
            }
        )
        
    except ImportError:
        raise HTTPException(
            status_code=500, 
            detail="python-docx library not installed. Please install it: pip install python-docx"
        )
    except Exception as e:
        print(f"❌ DOCX conversion error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to convert notes to DOCX: {str(e)}")


def add_formatted_text(paragraph, text):
    """Add text with markdown formatting to a paragraph"""
    # Simple markdown parsing for bold and italic
    # Handle **bold** and *italic*
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 1:
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            paragraph.add_run(part)


@app.post("/api/context/query")
async def api_query_context(request: Request, payload: dict = Body(...)):
    """Query operational context from Sindh Police documents"""
    token = get_token_from_request(request)
    verify_jwt_token(token)
    
    query = payload.get("query", "")
    if not query:
        raise HTTPException(status_code=400, detail="Query is required")
    
    result = request_regulatory_context(query)
    return result


@app.post("/api/transcript/store")
async def store_transcript_entry(request: Request, payload: dict = Body(...)):
    """Store a transcript entry from frontend"""
    token = get_token_from_request(request)
    verify_jwt_token(token)
    
    meeting_id = payload.get("meeting_id")
    speaker = payload.get("speaker")
    text = payload.get("text")
    
    if not all([meeting_id, speaker, text]):
        raise HTTPException(status_code=400, detail="meeting_id, speaker, and text are required")
    
    result = add_transcript_entry(meeting_id, speaker, text)
    return result


# =============================================================================
# HEYGEN AVATAR INTEGRATION ENDPOINTS - REMOVED
# =============================================================================

# =============================================================================
# BROWSER CALL / VOICE ENDPOINTS
# =============================================================================

# =============================================================================
# BROWSER CALL / VOICE ENDPOINTS
# =============================================================================

@app.post("/start-browser-call")
async def start_browser_call(request: Request, payload: dict = Body(...)):
    """Start a browser-based voice session for the meeting"""
    global active_meeting_id
    
    token = get_token_from_request(request)
    user_data = verify_jwt_token(token)
    
    if user_data.get("role") not in ["secretary", "admin"]:
        raise HTTPException(status_code=403, detail="Only Secretary or Admin can start voice sessions")
    
    meeting_id = payload.get("meeting_id", active_meeting_id)
    
    if not meeting_id:
        # Auto-start a meeting if none exists
        meeting_id = f"MEETING-{dt.now().strftime('%Y%m%d%H%M%S')}"
        start_meeting_session(meeting_id)
        active_meeting_id = meeting_id
    
    call_id = str(uuid.uuid4())
    call_recordings[call_id] = {"incoming": [], "outgoing": [], "start_time": time.time()}
    
    meeting_metadata[call_id] = {
        "meeting_id": meeting_id,
        "user": user_data["username"],
        "role": user_data["role"]
    }
    
    print(f"🎙️ Voice session started for meeting: {meeting_id}")
    
    return {
        "call_id": call_id,
        "meeting_id": meeting_id,
        "voice": VOICE
    }


# WebSocket for browser audio streaming
import websockets as ws_client

@app.websocket("/media-stream-browser")
async def media_stream_browser(websocket: WebSocket):
    """WebSocket endpoint for real-time audio streaming"""
    await websocket.accept()

    openai_url = 'wss://api.openai.com/v1/realtime?model=gpt-4o-realtime-preview-2025-06-03'
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "OpenAI-Beta": "realtime=v1"
    }
  
    async with ws_client.connect(openai_url, additional_headers=headers) as openai_ws:
        session_initialized = False
        call_id = None
        meeting_id = None

        user_pcm_buffer = io.BytesIO()
        agent_pcm_buffer = io.BytesIO()

        async def receive_from_browser():
            nonlocal session_initialized, call_id, meeting_id
            
            async for msg in websocket.iter_text():
                data = json.loads(msg)

                if data.get("event") == "start":
                    # Verify JWT token
                    token = data["start"]["customParameters"].get("token")
                    if not token:
                        print("❌ No token provided")
                        await websocket.close(code=1008, reason="Authentication required")
                        return
                    
                    try:
                        user_data = verify_jwt_token(token)
                        print(f"✅ WebSocket authenticated: {user_data['username']} ({user_data['role']})")
                    except HTTPException as e:
                        print(f"❌ Invalid token: {e.detail}")
                        await websocket.close(code=1008, reason="Invalid token")
                        return
                    
                    call_id = data["start"]["customParameters"].get("call_id")
                    meeting_id = data["start"]["customParameters"].get("meeting_id", active_meeting_id)
                    
                    # Initialize OpenAI session with Sindh Police context
                    await initialize_session(openai_ws, call_id, meeting_id)
                    await send_initial_conversation_item(openai_ws)
                    session_initialized = True
                    continue

                if data.get("event") == "media" and session_initialized:
                    payload_b64 = data["media"]["payload"]
                    pcm_bytes = base64.b64decode(payload_b64)
                    user_pcm_buffer.write(pcm_bytes)
                    
                    # Send audio to OpenAI
                    mulaw_bytes = audioop.lin2ulaw(pcm_bytes, 2)
                    audio_append = {
                        "type": "input_audio_buffer.append",
                        "audio": base64.b64encode(mulaw_bytes).decode('utf-8')
                    }
                    await openai_ws.send(json.dumps(audio_append))

                # Handle motion submission for voting
                if data.get("event") == "motion" and session_initialized:
                    motion_text = data.get("motion_text", "")
                    regulatory_context = data.get("regulatory_context", "")
                    
                    print(f"📋 Motion received for voting: {motion_text[:50]}...")
                    
                    # Send motion to AI as a text message for voting
                    motion_message = {
                        "type": "conversation.item.create",
                        "item": {
                            "type": "message",
                            "role": "user",
                            "content": [
                                {
                                    "type": "input_text",
                                    "text": f"""VOTING ITEM SUBMITTED:

Motion: {motion_text}

Relevant Operational Context:
{regulatory_context if regulatory_context else 'No specific context found.'}

Please analyze this motion against Sindh Police policies and operational realities and cast your vote using the cast_vote function. Provide your vote (FOR/AGAINST/ABSTAIN), reasoning, policy reference, and risk assessment."""
                                }
                            ]
                        }
                    }
                    await openai_ws.send(json.dumps(motion_message))
                    await openai_ws.send(json.dumps({"type": "response.create"}))
                    continue

                if data.get("event") == "stop":
                    break

        async def receive_from_openai_and_forward():
            # Track current response transcript to detect acknowledgment phrases
            current_response_text = ""
            suppress_audio = False
            
            # Phrases to filter out (case-insensitive, with variations)
            FILTER_PHRASES = [
                "listening silently",
                "Listening silently",
                "sun rahi hun",
                 "Sun r ahi hun",
                "sun rahi hoon",
                "sun rahi hoon.",
                "sun rahi hun.",
                "i'm listening",
                "im listening"
            ]
            
            # Buffer for accumulating user transcript deltas
            user_transcript_buffer = ""
            
            # Buffer for accumulating function call arguments
            function_args_buffer = ""
            current_function_name = None
            
            async for raw in openai_ws:
                response = json.loads(raw)
                rtype = response.get("type")

                if rtype == 'input_audio_buffer.speech_started':
                    print("🎤 Speech detected - interruption")
                    await openai_ws.send(json.dumps({"type": "response.cancel"}))
                    await websocket.send_json({"event": "clear"})
                    # Reset tracking
                    current_response_text = ""
                    suppress_audio = False
                    continue
                
                if rtype in LOG_EVENT_TYPES:
                    continue
                
                # Debug: Log all conversation item events to understand what we're receiving
                if rtype and "conversation" in rtype.lower():
                    print(f"🔍 Conversation event: {rtype}")

                # Reset tracking when response starts
                if rtype == "response.created":
                    current_response_text = ""
                    suppress_audio = False
                    # Also reset function call buffers
                    function_args_buffer = ""
                    current_function_name = None
                
                # Track transcript to detect acknowledgment phrases
                if rtype == "response.audio_transcript.delta":
                    transcript_delta = response.get("delta", "")
                    current_response_text += transcript_delta
                    
                    # Check if current text matches any filter phrase
                    text_lower = current_response_text.lower().strip()
                    for phrase in FILTER_PHRASES:
                        if phrase in text_lower:
                            suppress_audio = True
                            print(f"🔇 Filtering acknowledgment phrase: {current_response_text.strip()}")
                            break
                    
                    # Forward transcript to frontend only if not filtered
                    if transcript_delta and meeting_id and not suppress_audio:
                        print(f"🚀 SENDING to frontend: transcript event [Sindh Police AI]: {transcript_delta[:50]}")
                        await websocket.send_json({
                            "event": "transcript",
                            "speaker": "Sindh Police AI",
                            "text": transcript_delta
                        })
                        print(f"✅ SENT to frontend successfully")
                        # Store transcript entry in meeting session
                        if meeting_id:
                            add_transcript_entry(meeting_id, "Sindh Police AI", transcript_delta)
                
                # Handle user speech transcript from conversation items
                # OpenAI Realtime API provides user transcripts through multiple events
                
                # Event 1: Direct transcription completed event
                if rtype == "conversation.item.input_audio_transcription.completed":
                    user_transcript = response.get("transcript", "")
                    if user_transcript and meeting_id:
                        print(f"📝 User transcript (completed): {user_transcript[:100]}...")
                        add_transcript_entry(meeting_id, "User", user_transcript)
                        print(f"🚀 SENDING to frontend: transcript event [User]: {user_transcript[:50]}")
                        await websocket.send_json({
                            "event": "transcript",
                            "speaker": "User",
                            "text": user_transcript
                        })
                        print(f"✅ SENT to frontend successfully")
                
                # Event 2: Transcription delta events (word-by-word accumulation)
                if rtype == "conversation.item.input_audio_transcription.delta":
                    delta = response.get("delta", "")
                    if delta:
                        user_transcript_buffer += delta
                        # Optionally send incremental updates to frontend
                        if meeting_id:
                            await websocket.send_json({
                                "event": "transcript",
                                "speaker": "User",
                                "text": delta,
                                "incremental": True
                            })
                
                # Event 3: Transcription done (final complete transcript)
                if rtype == "conversation.item.input_audio_transcription.done":
                    # Get the final transcript from the item
                    item = response.get("item", {})
                    final_transcript = item.get("transcript", "") or user_transcript_buffer
                    
                    if final_transcript and meeting_id:
                        print(f"📝 User transcript (done): {final_transcript[:100]}...")
                        add_transcript_entry(meeting_id, "User", final_transcript)
                        await websocket.send_json({
                            "event": "transcript",
                            "speaker": "User",
                            "text": final_transcript
                        })
                    # Clear buffer
                    user_transcript_buffer = ""
                
                # Event 4: Conversation item created with transcription
                if rtype == "conversation.item.created":
                    item = response.get("item", {})
                    
                    # Check if it's a user message
                    if item.get("type") == "message" and item.get("role") == "user":
                        # Check content for transcription
                        content = item.get("content", [])
                        for content_item in content:
                            # Handle input_audio_transcription type
                            if content_item.get("type") == "input_audio_transcription":
                                transcript_text = content_item.get("transcript", "")
                                if transcript_text and meeting_id:
                                    print(f"📝 User transcript (item.created): {transcript_text[:100]}...")
                                    add_transcript_entry(meeting_id, "User", transcript_text)
                                    await websocket.send_json({
                                        "event": "transcript",
                                        "speaker": "User",
                                        "text": transcript_text
                                    })
                            # Handle input_text type (for motions or typed input)
                            elif content_item.get("type") == "input_text":
                                text_content = content_item.get("text", "")
                                # Only store if it's not a voting item (those are handled separately)
                                if text_content and meeting_id and "VOTING ITEM SUBMITTED" not in text_content:
                                    print(f"📝 User text input: {text_content[:100]}...")
                                    add_transcript_entry(meeting_id, "User", text_content)
                                    await websocket.send_json({
                                        "event": "transcript",
                                        "speaker": "User",
                                        "text": text_content
                                    })
                
                # Event 5: Input audio buffer committed (user finished speaking)
                if rtype == "input_audio_buffer.committed":
                    print("🎤 User speech committed to buffer - waiting for transcription...")
                    # The transcript will come through conversation.item events above
                
                # Reset tracking when response is done
                if rtype == "response.done":
                    current_response_text = ""
                    suppress_audio = False

                if rtype == "response.audio.delta" and "delta" in response:
                    # Skip audio if this is an acknowledgment phrase
                    if suppress_audio:
                        continue
                    
                    mulaw_b64 = response["delta"]
                    mulaw_bytes = base64.b64decode(mulaw_b64)

                    try:
                        pcm = audioop.ulaw2lin(mulaw_bytes, 2)
                    except Exception:
                        pcm = mulaw_bytes
                    
                    agent_pcm_buffer.write(pcm)
                    pcm_b64 = base64.b64encode(pcm).decode('utf-8')

                    out = {
                        "event": "media",
                        "media": {
                            "payload": pcm_b64,
                            "format": "raw_pcm",
                            "sampleRate": 8000,
                            "channels": 1,
                            "bitDepth": 16
                        }
                    }
                    await websocket.send_json(out)

                # Handle function call arguments (accumulate deltas)
                if rtype == "response.function_call_arguments.delta":
                    delta = response.get("delta", "")
                    function_args_buffer += delta
                    current_function_name = response.get("name", current_function_name)
                
                # Handle function calls (voting, etc.)
                elif rtype == "response.function_call_arguments.done":
                    func_name = response.get("name", current_function_name)
                    # Get complete arguments from response or buffer
                    arguments_str = response.get("arguments", function_args_buffer)
                    
                    # Try to parse JSON with error handling
                    try:
                        func_args = json.loads(arguments_str) if arguments_str else {}
                    except json.JSONDecodeError as e:
                        print(f"⚠️ JSON decode error for function arguments: {e}")
                        print(f"   Arguments string: {arguments_str[:200]}...")
                        # Try to fix common JSON issues
                        try:
                            # Try to complete the JSON if it's truncated
                            if arguments_str and not arguments_str.strip().endswith('}'):
                                # Attempt to close the JSON
                                fixed_args = arguments_str.rstrip()
                                # Remove trailing comma if present
                                if fixed_args.endswith(','):
                                    fixed_args = fixed_args[:-1]
                                # Try to close the object
                                if not fixed_args.endswith('}'):
                                    fixed_args += '}'
                                func_args = json.loads(fixed_args)
                            else:
                                func_args = {}
                        except:
                            print(f"❌ Could not parse function arguments, using empty dict")
                            func_args = {}
                    
                    print(f"🔧 Function call: {func_name}")
                    
                    # Process the function call
                    result = await handle_function_call(func_name, func_args, meeting_id)
                    
                    # Log the result for debugging
                    print(f"📊 Function result for {func_name}: {json.dumps(result, indent=2)[:500]}")
                    
                    # Send result back to frontend
                    outgoing = {
                        "event": "function_result", 
                        "name": func_name,
                        "arguments": arguments_str,
                        "result": result
                    }
                    await websocket.send_json(outgoing)
                    print(f"📤 Sent function_result to frontend for {func_name}")
                    
                    # Send function result back to OpenAI for continuation
                    function_output = {
                        "type": "conversation.item.create",
                        "item": {
                            "type": "function_call_output",
                            "call_id": response.get("call_id"),
                            "output": json.dumps(result)
                        }
                    }
                    await openai_ws.send(json.dumps(function_output))
                    await openai_ws.send(json.dumps({"type": "response.create"}))
                    
                    # Reset buffers
                    function_args_buffer = ""
                    current_function_name = None
                
                # Reset function buffers when response starts
                if rtype == "response.created":
                    function_args_buffer = ""
                    current_function_name = None

        recv_task = asyncio.create_task(receive_from_browser())
        send_task = asyncio.create_task(receive_from_openai_and_forward())

        try:
            await recv_task
        finally:
            if not send_task.done():
                send_task.cancel()
            await websocket.close()


async def handle_function_call(func_name: str, func_args: dict, meeting_id: str) -> dict:
    """Handle function calls from the AI"""
    
    if func_name == "cast_vote":
        return cast_vote(
            meeting_id=meeting_id or "ADHOC",
            motion_description=func_args.get("motion_description", ""),
            vote=func_args.get("vote", "ABSTAIN"),
            reasoning=func_args.get("reasoning", ""),
            regulatory_reference=func_args.get("regulatory_reference", ""),
            risk_assessment=func_args.get("risk_assessment", "")
        )
    
    elif func_name == "request_clarification":
        return {
            "success": True,
            "topic": func_args.get("topic", ""),
            "question": func_args.get("question", ""),
            "message": "Clarification requested from the board"
        }
    
    elif func_name == "cite_regulation":
        # RAG feature removed - return without context
        return {
            "success": True,
            "document": func_args.get("document_name", ""),
            "section": func_args.get("section", ""),
            "relevance": func_args.get("relevance", ""),
            "context": "RAG feature has been removed."
        }
    
    return {"success": False, "error": f"Unknown function: {func_name}"}


async def initialize_session(openai_ws, call_id: str, meeting_id: str = None):
    """Initialize the OpenAI session with Sindh Police context"""
    
    # RAG feature removed - no regulatory context retrieval
    regulatory_context = ""
    
    # Build system message with context
    system_message = build_system_message(
        instructions="",
        caller="",
        voice=VOICE,
        regulatory_context=regulatory_context
    )
    
    print(f"🔧 Initializing Sindh Police AI Meeting Member session")

    session_update = {
        "type": "session.update",
        "session": {
            "turn_detection": {
                "type": "server_vad",
                "threshold": 0.8,
                "prefix_padding_ms": 500,
                "silence_duration_ms": 1500,
                "create_response": True,  # AI responds on each turn to acknowledge listening
                "interrupt_response": True,
            },
            "input_audio_format": "g711_ulaw",
            "output_audio_format": "g711_ulaw",
            "input_audio_transcription": {
                "model": "whisper-1"
            },
            "voice": VOICE,
            "instructions": system_message,
            "modalities": ["text", "audio"],
            "temperature": 0.7,
            "speed": 1.0,
            "tool_choice": "auto",
            "tools": function_call_tools
        }
    }
    
    await openai_ws.send(json.dumps(session_update))


async def send_initial_conversation_item(openai_ws):
    """Send initial context to the AI"""
    initial_message = {
        "type": "conversation.item.create",
        "item": {
            "type": "message",
            "role": "user",
            "content": [
                {
                    "type": "input_text",
                    "text": "Meeting started. Remember: Your unique name is 'Sindh Police AI'. Only provide full opinions when someone explicitly addresses you by this name (e.g., 'Sindh Police AI, what's your opinion?'). For all other discussions, just acknowledge with 'Listening silently' or 'Sun rahi hun'."
                }
            ]
        }
    }
    await openai_ws.send(json.dumps(initial_message))
    await openai_ws.send(json.dumps({"type": "response.create"}))


# =============================================================================
# HEALTH CHECK & UTILITY ENDPOINTS
# =============================================================================

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "service": "Sindh Police AI Meeting Member",
        "active_meeting": active_meeting_id,
        "timestamp": dt.now(timezone.utc).isoformat()
    }


@app.get("/api/system/status")
async def system_status(request: Request):
    """Get system status including AI connectivity"""
    token = get_token_from_request(request)
    verify_jwt_token(token)
    
    return {
        "status": "online",
        "ai_ready": True,
        "active_meeting": active_meeting_id,
        "total_meetings": len(meeting_sessions),
        "total_votes": len(vote_history)
    }


# =============================================================================
# DOCUMENT MANAGEMENT ENDPOINTS
# =============================================================================

# In-memory document tracking (for listing purposes)
uploaded_documents = []

@app.get("/admin", response_class=HTMLResponse)
async def admin_page():
    """Serve the Document Management page"""
    with open("static/admin.html", "r", encoding="utf-8") as f:
        html_content = f.read()
    return html_content


@app.post("/api/documents/upload")
async def upload_documents(request: Request, files: List[UploadFile] = File(...)):
    """Upload and process documents into Pinecone vector database"""
    token = get_token_from_request(request)
    user_data = verify_jwt_token(token)
    
    # Only admin and secretary can upload
    if user_data.get("role") not in ["secretary", "admin"]:
        raise HTTPException(status_code=403, detail="Admin or Secretary access required")
    
    try:
        # Import document processing libraries
        from pypdf import PdfReader
        from docx import Document
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from langchain_openai import OpenAIEmbeddings
        from pinecone import Pinecone
        
        # Initialize
        pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        index = pc.Index(os.getenv("INDEX_NAME", "sindh-police-docs"))
        embeddings = OpenAIEmbeddings(
            model="text-embedding-3-small",
            openai_api_key=os.getenv("OPENAI_API_KEY")
        )
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        
        total_chunks = 0
        processed_docs = []
        
        for file in files:
            content = await file.read()
            filename = file.filename
            text = ""
            
            # Extract text based on file type
            if filename.endswith('.pdf'):
                pdf_reader = PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
            elif filename.endswith('.docx'):
                doc = Document(io.BytesIO(content))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            else:
                continue
            
            if not text.strip():
                continue
            
            # Split into chunks
            chunks = text_splitter.split_text(text)
            
            # Generate embeddings and upsert to Pinecone
            vectors = []
            for i, chunk in enumerate(chunks):
                # Create unique ID
                chunk_id = hashlib.md5(f"{filename}_{i}_{chunk[:50]}".encode()).hexdigest()
                
                # Generate embedding
                embedding = embeddings.embed_query(chunk)
                
                vectors.append({
                    "id": chunk_id,
                    "values": embedding,
                    "metadata": {
                        "text": chunk,
                        "source": filename,
                        "chunk_index": i,
                        "uploaded_at": dt.now(timezone.utc).isoformat(),
                        "uploaded_by": user_data["username"]
                    }
                })
            
            # Batch upsert (Pinecone limit is 100)
            for i in range(0, len(vectors), 100):
                batch = vectors[i:i+100]
                index.upsert(vectors=batch)
            
            total_chunks += len(chunks)
            
            # Track document
            doc_info = {
                "name": filename,
                "chunks": len(chunks),
                "uploaded_at": dt.now().strftime("%Y-%m-%d %H:%M"),
                "uploaded_by": user_data["username"]
            }
            processed_docs.append(doc_info)
            uploaded_documents.append(doc_info)
            
            print(f"✅ Processed {filename}: {len(chunks)} chunks")
        
        return {
            "success": True,
            "documents_processed": len(processed_docs),
            "chunks_created": total_chunks,
            "documents": processed_docs
        }
        
    except Exception as e:
        print(f"❌ Document upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/documents/list")
async def list_documents(request: Request):
    """List all documents in the knowledge base by querying Pinecone directly"""
    global uploaded_documents
    
    token = get_token_from_request(request)
    user_data = verify_jwt_token(token)
    
    if user_data.get("role") not in ["secretary", "admin"]:
        raise HTTPException(status_code=403, detail="Admin or Secretary access required")
    
    try:
        from pinecone import Pinecone
        
        pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        index = pc.Index(os.getenv("INDEX_NAME", "sindh-police-docs"))
        
        # Get index stats
        stats = index.describe_index_stats()
        total_vectors = stats.total_vector_count
        
        # Query Pinecone to discover unique documents from metadata
        document_map = {}
        
        if total_vectors > 0:
            try:
                embedding_dim = 1536  # text-embedding-3-small dimension
                dummy_vector = [0.0] * embedding_dim
                
                # Sample vectors to discover unique document sources
                sample_size = min(1000, total_vectors)
                query_result = index.query(
                    vector=dummy_vector,
                    top_k=sample_size,
                    include_metadata=True
                )
                
                # Extract unique document sources from metadata
                unique_sources = set()
                for match in query_result.matches:
                    metadata = match.metadata or {}
                    source = metadata.get("source")
                    if source:
                        unique_sources.add(source)
                
                # For each unique document, get accurate chunk count and metadata
                for doc_name in unique_sources:
                    try:
                        # Query with filter to count chunks for this document
                        count_query = index.query(
                            vector=dummy_vector,
                            top_k=10000,  # Max allowed by Pinecone
                            filter={"source": doc_name},
                            include_metadata=True
                        )
                        
                        if count_query.matches:
                            chunk_count = len(count_query.matches)
                            # Get metadata from first match
                            first_match = count_query.matches[0]
                            metadata = first_match.metadata or {}
                            
                            document_map[doc_name] = {
                                "name": doc_name,
                                "chunks": chunk_count,
                                "uploaded_at": metadata.get("uploaded_at", "Unknown"),
                                "uploaded_by": metadata.get("uploaded_by", "Unknown")
                            }
                    except Exception as e:
                        print(f"⚠️ Error querying for {doc_name}: {e}")
            except Exception as e:
                print(f"⚠️ Error querying Pinecone: {e}")
        
        # Convert to list and sort by name
        documents_list = list(document_map.values())
        documents_list.sort(key=lambda x: x["name"])
        
        # Also update the in-memory list for consistency
        uploaded_documents = documents_list
        
        return {
            "success": True,
            "total_documents": len(documents_list),
            "total_chunks": total_vectors,
            "documents": documents_list
        }
        
    except Exception as e:
        print(f"❌ Document list error: {e}")
        import traceback
        traceback.print_exc()
        # Fallback to in-memory list
        return {
            "success": True,
            "total_documents": len(uploaded_documents),
            "total_chunks": 0,
            "documents": [{
                "name": doc["name"],
                "chunks": doc.get("chunks", 0),
                "uploaded_at": doc.get("uploaded_at", "Unknown"),
                "uploaded_by": doc.get("uploaded_by", "Unknown")
            } for doc in uploaded_documents]
        }


@app.delete("/api/documents/delete/{document_name:path}")
async def delete_document(document_name: str, request: Request):
    """Delete all chunks for a specific document from Pinecone"""
    global uploaded_documents
    
    token = get_token_from_request(request)
    user_data = verify_jwt_token(token)
    
    if user_data.get("role") not in ["secretary", "admin"]:
        raise HTTPException(status_code=403, detail="Admin or Secretary access required")
    
    try:
        from pinecone import Pinecone
        
        pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        index = pc.Index(os.getenv("INDEX_NAME", "sindh-police-docs"))
        
        # Get chunk count before deletion for reporting
        doc_info = next((doc for doc in uploaded_documents if doc["name"] == document_name), None)
        chunk_count = doc_info.get("chunks", 0) if doc_info else 0
        
        # Delete all vectors with this source using metadata filter
        index.delete(
            filter={"source": document_name}
        )
        
        # Remove from uploaded_documents list
        uploaded_documents = [doc for doc in uploaded_documents if doc["name"] != document_name]
        
        print(f"🗑️ Deleted document '{document_name}': {chunk_count} chunks removed")
        
        return {
            "success": True,
            "message": f"Document '{document_name}' deleted successfully",
            "chunks_deleted": chunk_count
        }
        
    except Exception as e:
        print(f"❌ Document delete error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")


# =============================================================================
# LIVE MEETING & VOICE RECORDING ENDPOINTS
# =============================================================================

@app.get("/record", response_class=HTMLResponse)
async def record_page():
    """Serve the Voice Recording Vote page"""
    with open("static/record.html", "r", encoding="utf-8") as f:
        html_content = f.read()
    return html_content


@app.post("/api/vote/record")
async def vote_from_recording(request: Request, audio: UploadFile = File(...)):
    """
    Process a voice recording:
    1. Transcribe using OpenAI Whisper
    2. Get regulatory context from Pinecone
    3. Analyze and cast vote using OpenAI
    """
    token = get_token_from_request(request)
    user_data = verify_jwt_token(token)
    
    try:
        import openai
        import tempfile
        
        # Read audio file
        audio_content = await audio.read()
        original_filename = audio.filename or "recording.webm"
        
        # Determine input format from filename
        input_format = original_filename.split('.')[-1].lower()
        print(f"🎤 Processing audio (format: {input_format})...")
        
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        transcription_response = None
        
        # Try direct upload first for supported formats
        supported_formats = ['mp3', 'mp4', 'm4a', 'wav', 'mpeg', 'mpga', 'ogg']
        
        if input_format in supported_formats:
            # Direct upload for natively supported formats
            print("🎤 Transcribing audio (direct)...")
            audio_buffer = io.BytesIO(audio_content)
            audio_buffer.name = original_filename
            
            transcription_response = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_buffer,
                response_format="text"
            )
        else:
            # Convert webm/other formats to mp3 using pydub
            print("🔄 Converting audio format...")
            
            # Save to temp file
            suffix = f".{input_format}"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_input:
                temp_input.write(audio_content)
                temp_input_path = temp_input.name
            
            try:
                # Convert using pydub (requires ffmpeg)
                audio_segment = AudioSegment.from_file(temp_input_path)
                
                # Export to mp3
                temp_mp3_path = temp_input_path.rsplit('.', 1)[0] + ".mp3"
                audio_segment.export(temp_mp3_path, format="mp3")
                
                print("🎤 Transcribing audio...")
                with open(temp_mp3_path, "rb") as audio_file:
                    transcription_response = client.audio.transcriptions.create(
                        model="whisper-1",
                        file=audio_file,
                        response_format="text"
                    )
                
                # Clean up temp files
                os.unlink(temp_input_path)
                os.unlink(temp_mp3_path)
                
            except Exception as conv_error:
                print(f"⚠️ Conversion failed: {conv_error}")
                # Clean up on error
                if os.path.exists(temp_input_path):
                    os.unlink(temp_input_path)
                
                # Last resort: try direct with original name
                audio_buffer = io.BytesIO(audio_content)
                audio_buffer.name = original_filename
                
                transcription_response = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_buffer,
                    response_format="text"
                )
        
        transcription = transcription_response.strip()
        print(f"📝 Transcription: {transcription[:100]}...")
        
        if not transcription:
            raise HTTPException(status_code=400, detail="Could not transcribe audio")
        
        # Step 2: Analyze and vote using OpenAI (RAG feature removed)
        print("🤖 Analyzing motion and casting vote...")
        
        analysis_prompt = f"""You are the Sindh Police AI Meeting Member. Analyze the following motion/proposal and cast your vote.

MOTION/PROPOSAL (transcribed from voice recording):
{transcription}

Based on Sindh Police policies and operational realities, cast your vote on this motion.

You must respond in the following JSON format ONLY (no other text):
{{
    "motion_summary": "A concise 1-2 sentence summary of the motion/agenda being proposed",
    "vote": "FOR" or "AGAINST" or "ABSTAIN",
    "reasoning": "Clear explanation of your decision (2-3 sentences)",
    "regulatory_reference": "Specific Sindh Police policy/procedure supporting your vote",
    "risk_assessment": "Brief note on potential implications"
}}

VOTING GUIDELINES:
- Vote FOR if the motion aligns with Sindh Police mission and enhances public safety
- Vote AGAINST if the motion contradicts policies or compromises officer/public safety
- Vote ABSTAIN if there's insufficient information to decide

Respond with valid JSON only."""

        vote_response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are the Sindh Police AI Meeting Member. Respond only with valid JSON."},
                {"role": "user", "content": analysis_prompt}
            ],
            temperature=0.7,
            max_tokens=500
        )
        
        vote_text = vote_response.choices[0].message.content.strip()
        
        # Parse JSON response
        # Clean up response if it has markdown code blocks
        if vote_text.startswith("```"):
            vote_text = vote_text.split("```")[1]
            if vote_text.startswith("json"):
                vote_text = vote_text[4:]
        vote_text = vote_text.strip()
        
        vote_data = json.loads(vote_text)
        
        print(f"🗳️ Vote: {vote_data.get('vote')}")
        
        # Store the vote
        from tools import cast_vote as store_vote
        store_vote(
            meeting_id="RECORDING-" + dt.now().strftime('%Y%m%d%H%M%S'),
            motion_description=transcription[:200],
            vote=vote_data.get("vote", "ABSTAIN"),
            reasoning=vote_data.get("reasoning", ""),
            regulatory_reference=vote_data.get("regulatory_reference", ""),
            risk_assessment=vote_data.get("risk_assessment", "")
        )
        
        return {
            "success": True,
            "transcription": transcription,
            "vote": vote_data,
            "regulatory_context_used": True
        }
        
    except json.JSONDecodeError as e:
        print(f"❌ JSON parse error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse AI response")
    except Exception as e:
        print(f"❌ Recording vote error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    print("🚀 Starting Sindh Police AI Meeting Member...")
    print(f"📍 Server running at http://localhost:{PORT}")
    uvicorn.run(app, host="0.0.0.0", port=PORT)
